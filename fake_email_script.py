import os
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import base64
from urllib.parse import urlparse
import requests
from email import encoders
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.image import MIMEImage
import mimetypes
from datetime import datetime
import uuid

def extract_images_from_word(word_file_path):
    """
    從Word文檔中提取圖片，支持內嵌圖片和外部引用圖片
    """
    doc = Document(word_file_path)
    images = []
    
    # 獲取文檔關係
    rels = doc.part.rels
    
    for rel in rels.values():
        # 只處理圖片類型的關係
        if rel.reltype != RT.IMAGE:
            continue
            
        try:
            # 嘗試獲取內嵌圖片
            if hasattr(rel, 'target_part') and rel.target_part:
                image_bytes = rel.target_part.blob
                # 生成唯一的Content-ID
                content_id = str(uuid.uuid4())
                image_info = {
                    'bytes': image_bytes,
                    'filename': os.path.basename(rel.target_ref),
                    'content_id': content_id
                }
                images.append(image_info)
            
            # 處理外部引用圖片
            elif rel.target_mode == 'External':
                target_url = rel.target_ref
                content_id = str(uuid.uuid4())
                # 檢查是否是本地文件路徑
                if os.path.exists(target_url):
                    with open(target_url, 'rb') as f:
                        image_bytes = f.read()
                        image_info = {
                            'bytes': image_bytes,
                            'filename': os.path.basename(target_url),
                            'content_id': content_id
                        }
                        images.append(image_info)
                # 檢查是否是網絡URL
                elif urlparse(target_url).scheme in ['http', 'https']:
                    response = requests.get(target_url)
                    if response.status_code == 200:
                        image_info = {
                            'bytes': response.content,
                            'filename': os.path.basename(urlparse(target_url).path) or 'image.jpg',
                            'content_id': content_id
                        }
                        images.append(image_info)
                    else:
                        print(f"警告：無法下載圖片 {target_url}")
                else:
                    print(f"警告：無法處理的圖片引用 {target_url}")
            
        except Exception as e:
            print(f"警告：處理圖片時發生錯誤: {str(e)}")
            continue
    
    return images

def has_inline_image(paragraph):
    """
    檢查段落是否包含內嵌圖片
    """
    for run in paragraph.runs:
        if 'graphicData' in str(run._element.xml) or 'picture' in str(run._element.xml):
            return True
    return False

def extract_content_from_word(word_file_path):
    """
    從Word文檔中提取文本和圖片位置
    """
    doc = Document(word_file_path)
    content_parts = []
    image_index = 0
    
    for paragraph in doc.paragraphs:
        if has_inline_image(paragraph):
            content_parts.append(f'<img src="cid:image_{image_index}" style="max-width:100%;"/>')
            image_index += 1
        elif paragraph.text.strip():
            # 處理段落樣式
            style = ''
            if paragraph.style.name.startswith('Heading'):
                level = paragraph.style.name[-1]
                content_parts.append(f'<h{level}>{paragraph.text}</h{level}>')
            else:
                content_parts.append(f'<p>{paragraph.text}</p>')
    
    return '\n'.join(content_parts)

def create_eml_with_content(word_file_path):
    """
    創建包含Word內容的EML文件，圖片以內嵌方式顯示
    """
    # 創建郵件對象
    msg = MIMEMultipart('related')
    
    # 設置基本郵件頭
    msg['Subject'] = 'Document Content'
    msg['From'] = 'sender@example.com'
    msg['To'] = 'recipient@example.com'
    msg['Date'] = datetime.now().strftime('%a, %d %b %Y %H:%M:%S +0000')
    
    # 創建HTML內容部分
    html_content = extract_content_from_word(word_file_path)
    html_part = MIMEText(f'''
    <html>
    <head>
        <style>
            body {{ font-family: Arial, sans-serif; line-height: 1.6; }}
            p {{ margin: 1em 0; }}
            img {{ display: block; margin: 1em 0; }}
        </style>
    </head>
    <body>
    {html_content}
    </body>
    </html>
    ''', 'html', 'utf-8')
    
    msg.attach(html_part)
    
    # 提取並添加內嵌圖片
    images = extract_images_from_word(word_file_path)
    for idx, image_info in enumerate(images):
        image_bytes = image_info['bytes']
        
        # 創建圖片部分
        image = MIMEImage(image_bytes)
        image.add_header('Content-ID', f'<image_{idx}>')
        image.add_header('Content-Disposition', 'inline')
        msg.attach(image)
    
    # 生成EML文件
    output_path = os.path.splitext(word_file_path)[0] + '.eml'
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(msg.as_string())
    
    print(f"已創建EML文件：{output_path}")
    return output_path

def main():
    """
    主函數
    """
    try:
        # 獲取當前腳本所在目錄
        script_dir = os.path.dirname(os.path.abspath(__file__))
        
        # 指定Word文件名
        word_filename = "test.docx"  # 替換為你的Word文件名
        
        # 構建完整的文件路徑
        word_file_path = os.path.join(script_dir, word_filename)
        
        # 檢查文件是否存在
        if not os.path.exists(word_file_path):
            print(f"錯誤：找不到Word文件 {word_file_path}")
            return
        
        # 創建EML文件
        output_path = create_eml_with_content(word_file_path)
        print(f"成功處理完成！EML文件保存在：{output_path}")
        
    except Exception as e:
        print(f"處理過程中發生錯誤：{str(e)}")

if __name__ == "__main__":
    main()